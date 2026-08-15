from __future__ import annotations

import fitz  # PyMuPDF
import docx  # python-docx

def extract_text_from_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    pages = []
    try:
        for page in doc:
            text = page.get_text().strip()
            if text:
                pages.append(text)
    finally:
        doc.close()
    return "\n\n".join(pages)


def extract_text_from_docx(file_path: str) -> str:

    document = docx.Document(file_path)
    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def extract_text_from_txt(file_path: str) -> str:
    """Plain text / markdown files -- just read them directly."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


_VALID_DOC_TYPES = {"resume", "technical_report", "general"}
_CLASSIFY_EXCERPT_CHARS = 2000


def classify_doc_type(text: str, source: str = "") -> str:
    """
    Fast and accurate classifier for document types ('resume', 'technical_report', 'general')
    based on filename keywords and document content structure. Runs in <1ms to prevent
    upload latency bottlenecks.
    """
    source_lower = source.lower()
    text_lower = (text or "")[:_CLASSIFY_EXCERPT_CHARS].lower()

    # 1. Fast heuristics based on filename
    if any(k in source_lower for k in ("cv", "resume", "curriculum_vitae", "curriculum vitae", "profile")):
        print(f"[classify_doc_type] filename match for '{source}' -> 'resume'")
        return "resume"
    if any(k in source_lower for k in ("tutorial", "report", "datasheet", "manual", "lab_", "lab-", "architecture", "guide", "lecture")):
        print(f"[classify_doc_type] filename match for '{source}' -> 'technical_report'")
        return "technical_report"

    # 2. Heuristics based on document content
    resume_keywords = (
        "education", "experience", "skills", "projects", "professional summary",
        "work experience", "selected projects", "technical skills", "linkedin.com",
        "github.com", "curriculum vitae", "undergraduate", "bachelor"
    )
    resume_hits = sum(1 for kw in resume_keywords if kw in text_lower)
    if resume_hits >= 3:
        print(f"[classify_doc_type] content score={resume_hits} for '{source}' -> 'resume'")
        return "resume"

    tech_keywords = (
        "abstract", "introduction", "methodology", "architecture", "implementation",
        "schematic", "registers", "simulation", "circuit", "protocol", "overview", "chapter"
    )
    tech_hits = sum(1 for kw in tech_keywords if kw in text_lower)
    if tech_hits >= 3:
        print(f"[classify_doc_type] content score={tech_hits} for '{source}' -> 'technical_report'")
        return "technical_report"

    return "general"